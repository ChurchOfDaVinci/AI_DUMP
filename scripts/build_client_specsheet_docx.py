"""
build_client_specsheet_docx.py
==============================
Converts Documents/400A_Charging_System_Client_Specsheet.md into a client-facing
Word (.docx) document that inherits styles, header/footer, and page setup from a
project-plan template document.

Usage (from the repository root):
    python scripts/build_client_specsheet_docx.py

Output:
    ACDC charger/400A Charging System Client Specsheet.docx

Requirements:
    pip install python-docx
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
REPO_ROOT = Path(__file__).resolve().parent.parent
TEMPLATE = REPO_ROOT / "ACDC charger" / "ProjectPlan_400A_Charger.docx"
SOURCE_MD = REPO_ROOT / "Documents" / "400A_Charging_System_Client_Specsheet.md"
OUTPUT = REPO_ROOT / "ACDC charger" / "400A Charging System Client Specsheet.docx"
FORBIDDEN_TERMS = ("UR100040", "UUGreenPower")

# ---------------------------------------------------------------------------
# Colour palette (matches the template's theme accent1 = 156082 teal, so the
# generated document keeps the same colour scheme the client applied by hand)
# ---------------------------------------------------------------------------
ACCENT = "156082"          # section-group header shading (theme accent1)
ACCENT_DARK = "0F4761"     # title / heading text (accent1 darker 25%)
ROW_STRIPE = "E7EEF2"      # light teal zebra striping for readability
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


# ---------------------------------------------------------------------------
# Helper: clear body content while keeping section properties (page/margin
# setup, header/footer links)
# ---------------------------------------------------------------------------
def _clear_body(doc: Document) -> None:
    """Remove all paragraphs and tables from the body, preserving sectPr."""
    body = doc.element.body
    to_remove = [child for child in body if child.tag != qn("w:sectPr")]
    for child in to_remove:
        body.remove(child)


# ---------------------------------------------------------------------------
# Helper: add a paragraph with optional bold/italic and inline markdown marks
# ---------------------------------------------------------------------------
def _add_para(doc: Document, text: str, style: str = "Normal", bold: bool = False, italic: bool = False) -> None:
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
            if bold:
                run.bold = True
            if italic:
                run.italic = True


# ---------------------------------------------------------------------------
# Helper: add a bullet list item
# ---------------------------------------------------------------------------
def _add_bullet(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="List Paragraph")
    p_pr = p._p.get_or_add_pPr()
    num_pr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    num_id = OxmlElement("w:numId")
    num_id.set(qn("w:val"), "1")
    num_pr.append(ilvl)
    num_pr.append(num_id)
    p_pr.append(num_pr)
    p.add_run(text)


# ---------------------------------------------------------------------------
# Helper: add a horizontal rule (paragraph with bottom border)
# ---------------------------------------------------------------------------
def _add_hr(doc: Document) -> None:
    p = doc.add_paragraph(style="Normal")
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "4472C4")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


# ---------------------------------------------------------------------------
# Low-level cell helpers
# ---------------------------------------------------------------------------
def _shade_cell(cell, fill: str) -> None:
    """Apply a solid background fill to a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def _set_cell_text(cell, text: str, bold: bool = False, italic: bool = False,
                   color: RGBColor | None = None, size: int | None = None) -> None:
    """Replace a cell's content with a single styled run."""
    p = cell.paragraphs[0]
    p.clear()
    # Parse inline **bold** / *italic* markers so markdown emphasis survives.
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    if not any(parts):
        parts = [text]
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
            run.bold = bold
            run.italic = italic
        if color is not None:
            run.font.color.rgb = color
        if size is not None:
            run.font.size = Pt(size)


def _set_col_widths(table, widths_cm: list[float]) -> None:
    """Force column widths (Word honours per-cell widths most reliably)."""
    table.autofit = False
    table.allow_autofit = False
    # Fixed layout + explicit grid so Word does not rebalance the columns.
    tbl_pr = table._tbl.tblPr
    layout = OxmlElement("w:tblLayout")
    layout.set(qn("w:type"), "fixed")
    tbl_pr.append(layout)

    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for col, width in zip(grid.findall(qn("w:gridCol")), widths_cm):
            col.set(qn("w:w"), str(int(width * 567)))  # 1 cm = 567 twips

    for row in table.rows:
        for cell, width in zip(row.cells, widths_cm):
            cell.width = Cm(width)


# ---------------------------------------------------------------------------
# Helper: build ONE consolidated specification table for the whole system.
# Each source section becomes a full-width shaded group header row, followed by
# its Parameter / Specification rows, giving a single clean datasheet table.
# ---------------------------------------------------------------------------
def _add_spec_table(doc: Document, blocks: list[dict]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Column header row (Parameter | Specification)
    head = table.add_row().cells
    for cell, label in zip(head, ("Parameter", "Specification")):
        _shade_cell(cell, ACCENT_DARK)
        _set_cell_text(cell, label, bold=True, color=WHITE)

    for index, block in enumerate(blocks, start=1):
        # Full-width group header row for the section.
        group_cells = table.add_row().cells
        merged = group_cells[0].merge(group_cells[1])
        _shade_cell(merged, ACCENT)
        _set_cell_text(merged, f"{index}.  {block['title']}", bold=True, color=WHITE)

        # Prose sections (e.g. System Overview) render as a full-width note row.
        for note in block.get("notes", []):
            note_cells = table.add_row().cells
            note_merged = note_cells[0].merge(note_cells[1])
            _set_cell_text(note_merged, note)

        # Parameter / specification rows with subtle zebra striping.
        for stripe, (param, spec) in enumerate(block.get("rows", [])):
            cells = table.add_row().cells
            _set_cell_text(cells[0], param, bold=True)
            _set_cell_text(cells[1], spec)
            if stripe % 2 == 1:
                _shade_cell(cells[0], ROW_STRIPE)
                _shade_cell(cells[1], ROW_STRIPE)

    _set_col_widths(table, [6.5, 10.5])
    doc.add_paragraph(style="Normal")


# ---------------------------------------------------------------------------
# Markdown parsing
# ---------------------------------------------------------------------------
def _split_markdown_cells(line: str) -> list[str]:
    return [cell.strip() for cell in line.strip().strip("|").split("|")]


def _parse_sections(md_text: str) -> tuple[list[dict], str]:
    lines = md_text.splitlines()
    sections: list[dict] = []
    disclaimer_lines: list[str] = []
    current: dict | None = None

    i = 0
    while i < len(lines):
        line = lines[i].strip()

        heading_match = re.match(r"^###\s+\d+\.\s+(.+)$", line)
        if heading_match:
            if current:
                sections.append(current)
            current = {"title": heading_match.group(1).strip(), "paragraphs": [], "table": None}
            i += 1
            continue

        if line.startswith("*All values"):
            disclaimer_lines = [line]
            i += 1
            while i < len(lines) and lines[i].strip():
                disclaimer_lines.append(lines[i].strip())
                i += 1
            break

        if current is None or line == "---" or not line:
            i += 1
            continue

        if line.startswith("|"):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].strip())
                i += 1

            if len(table_lines) >= 2:
                headers = _split_markdown_cells(table_lines[0])
                row_lines = [ln for ln in table_lines[2:] if ln.strip()]
                rows = [_split_markdown_cells(ln) for ln in row_lines]
                current["table"] = {"headers": headers, "rows": rows}
            continue

        current["paragraphs"].append(line)
        i += 1

    if current:
        sections.append(current)

    disclaimer = " ".join(disclaimer_lines).strip()
    if disclaimer.startswith("*") and disclaimer.endswith("*"):
        disclaimer = disclaimer[1:-1].strip()

    return sections, disclaimer


# ---------------------------------------------------------------------------
# Validation helpers
# ---------------------------------------------------------------------------
def _iter_document_text(doc: Document):
    for paragraph in doc.paragraphs:
        if paragraph.text:
            yield paragraph.text
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    yield cell.text


def _verify_no_forbidden_terms(docx_path: Path) -> None:
    doc = Document(str(docx_path))
    text_blob = "\n".join(_iter_document_text(doc)).lower()
    leaked = [term for term in FORBIDDEN_TERMS if term.lower() in text_blob]
    if leaked:
        raise ValueError(
            f"Forbidden internal terms found in generated document: {', '.join(leaked)}"
        )


# ---------------------------------------------------------------------------
# Build the document
# ---------------------------------------------------------------------------
def build(template_path: Path, source_md: Path, output_path: Path) -> None:
    md_text = source_md.read_text(encoding="utf-8")
    sections, disclaimer = _parse_sections(md_text)

    doc = Document(str(template_path))
    _clear_body(doc)

    # Title block --------------------------------------------------------
    title = doc.add_paragraph(style="Normal")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("400 A Liquid-Cooled DC Charging System")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string(ACCENT_DARK)

    subtitle = doc.add_paragraph(style="Normal")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("Technical Specification Sheet")
    sub_run.bold = True
    sub_run.font.size = Pt(14)
    sub_run.font.color.rgb = RGBColor.from_string(ACCENT)

    meta = doc.add_paragraph(style="Normal")
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta_run = meta.add_run(
        "Modular 280 kW integrated DC charging system  ·  Version 1.0  ·  2026-07-16"
    )
    meta_run.italic = True
    meta_run.font.size = Pt(10)

    _add_hr(doc)
    doc.add_paragraph(style="Normal")

    expected_headings = [
        "System Overview",
        "AC Input",
        "DC Output",
        "Performance & Efficiency",
        "Cooling System",
        "Control & Communication",
        "Environmental & Mechanical",
        "Compliance & Safety",
    ]

    by_title = {section["title"]: section for section in sections}

    blocks: list[dict] = []
    for heading in expected_headings:
        section = by_title.get(heading)
        if section is None:
            raise ValueError(f"Missing required section in markdown: {heading}")

        block: dict = {"title": heading, "notes": [], "rows": []}

        paragraphs = [p for p in section["paragraphs"] if p]
        if paragraphs:
            block["notes"].append(" ".join(paragraphs))

        table = section["table"]
        if table:
            block["rows"] = [(row[0], row[1]) for row in table["rows"] if len(row) >= 2]

        blocks.append(block)

    # One consolidated specification table for the whole system.
    _add_spec_table(doc, blocks)

    if disclaimer:
        _add_para(doc, disclaimer, style="Normal", italic=True)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))

    _verify_no_forbidden_terms(output_path)


if __name__ == "__main__":
    build(TEMPLATE, SOURCE_MD, OUTPUT)
