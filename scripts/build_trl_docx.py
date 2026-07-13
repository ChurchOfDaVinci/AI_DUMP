"""
build_trl_docx.py
=================
Converts Documents/TRL_Phases_Explained.md into a Word (.docx) document that
inherits the styles, header/footer, and page setup from the existing
"e-PU Cabinet V2 Measurement Plan.docx" template.

Usage (from the repository root):
    python scripts/build_trl_docx.py

Output:
    e-PU Cabinet V2/TRL Phases Explained.docx

Requirements:
    pip install python-docx
"""

from __future__ import annotations

import copy
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
REPO_ROOT = Path(__file__).resolve().parent.parent
TEMPLATE = REPO_ROOT / "e-PU Cabinet V2" / "e-PU Cabinet V2 Measurement Plan.docx"
SOURCE_MD = REPO_ROOT / "Documents" / "TRL_Phases_Explained.md"
OUTPUT = REPO_ROOT / "e-PU Cabinet V2" / "TRL Phases Explained.docx"


# ---------------------------------------------------------------------------
# Helper: clear body content while keeping section properties (page/margin
# setup, header/footer links)
# ---------------------------------------------------------------------------
def _clear_body(doc: Document) -> None:
    """Remove all paragraphs and tables from the body, preserving sectPr."""
    body = doc.element.body
    # Collect children to remove (everything except w:sectPr)
    to_remove = [
        child for child in body
        if child.tag != qn("w:sectPr")
    ]
    for child in to_remove:
        body.remove(child)


# ---------------------------------------------------------------------------
# Helper: add a paragraph with optional bold runs
# Supports inline **bold** markup in text.
# ---------------------------------------------------------------------------
def _add_para(doc: Document, text: str, style: str = "Normal",
               bold: bool = False, italic: bool = False) -> None:
    """Add a paragraph, supporting **bold** and *italic* inline markup."""
    p = doc.add_paragraph(style=style)

    # Split on **bold** and *italic* markers
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            run = p.add_run(part)
            if bold:
                run.bold = True
            if italic:
                run.italic = True


# ---------------------------------------------------------------------------
# Helper: add a bullet list item
# ---------------------------------------------------------------------------
def _add_bullet(doc: Document, text: str) -> None:
    """Add a bullet-list paragraph (List Paragraph style with bullet numPr)."""
    p = doc.add_paragraph(style="List Paragraph")
    # Apply list bullet numbering via direct XML (word needs numId / ilvl)
    pPr = p._p.get_or_add_pPr()
    numPr = OxmlElement("w:numPr")
    ilvl = OxmlElement("w:ilvl")
    ilvl.set(qn("w:val"), "0")
    numId = OxmlElement("w:numId")
    numId.set(qn("w:val"), "1")  # numId 1 = first list definition in template
    numPr.append(ilvl)
    numPr.append(numId)
    pPr.append(numPr)

    # Parse inline bold/italic
    parts = re.split(r"(\*\*[^*]+\*\*|\*[^*]+\*)", text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("*") and part.endswith("*"):
            run = p.add_run(part[1:-1])
            run.italic = True
        else:
            p.add_run(part)


# ---------------------------------------------------------------------------
# Helper: add a horizontal rule (paragraph with bottom border)
# ---------------------------------------------------------------------------
def _add_hr(doc: Document) -> None:
    p = doc.add_paragraph(style="Normal")
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "4472C4")
    pBdr.append(bottom)
    pPr.append(pBdr)


# ---------------------------------------------------------------------------
# Helper: add a table from a list of rows (list of list of strings)
# ---------------------------------------------------------------------------
def _add_table(doc: Document, headers: list[str], rows: list[list[str]]) -> None:
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row – bold
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.bold = True
        # Bold via direct run if text was set without runs
        if not hdr_cells[i].paragraphs[0].runs:
            hdr_cells[i].paragraphs[0].runs  # trigger
        p = hdr_cells[i].paragraphs[0]
        if p.runs:
            p.runs[0].bold = True
        else:
            run = p.add_run(h)
            run.bold = True
            # Remove the text set earlier
            p.clear()
            p.add_run(h).bold = True

    # Header row shading (blue, matching template heading color)
    for cell in hdr_cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "4472C4")
        tcPr.append(shd)
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                run.bold = True

    # Data rows
    for r_idx, row_data in enumerate(rows):
        row_cells = table.rows[r_idx + 1].cells
        for c_idx, val in enumerate(row_data):
            row_cells[c_idx].text = val

    doc.add_paragraph(style="Normal")  # spacing after table


# ---------------------------------------------------------------------------
# Build the document
# ---------------------------------------------------------------------------
def build(template_path: Path, source_md: Path, output_path: Path) -> None:
    doc = Document(str(template_path))
    _clear_body(doc)

    # ---- Title block -------------------------------------------------------
    _add_para(doc, "e-PU Cabinet V2", style="Normal")
    _add_para(doc, "TRL Phases Explained", style="Normal")
    doc.add_paragraph(style="Normal")
    _add_para(doc, "Technology Readiness Levels — General Reference Guide", style="Normal")
    doc.add_paragraph(style="Normal")
    _add_para(doc, "Version: 1.0", style="Normal")
    _add_para(doc, "Date: 2026-07-13", style="Normal")
    doc.add_paragraph(style="Normal")

    # ---- Introduction (Heading 1) ------------------------------------------
    doc.add_heading("Introduction", level=1)
    _add_para(doc,
        "This document explains what each Technology Readiness Level (TRL) phase means in "
        "general: a plain-language definition of the level, the goal of the phase, what is "
        "typically done in it, the environment/fidelity it is tested at, what is normally used to "
        "prove you have reached it, common pitfalls, and how the phase maps onto the projects in "
        "this repository. It is written as a general reference so it can be applied to any of the "
        "hardware/engineering projects here (e.g. the 400 A Charger, the e-PU Cabinet V2 BESS, "
        "etc.).",
        style="Normal")
    doc.add_paragraph(style="Normal")
    _add_para(doc,
        "What is TRL?  TRL is a 1–9 scale, originally developed by NASA and later adopted by "
        "the European Union (Horizon 2020 / Horizon Europe) and industry worldwide, to measure "
        "how mature a technology is — from a first idea (TRL 1) to a fully proven, commercially "
        "deployed product (TRL 9). Each step up means more evidence, more integration, and a "
        "more realistic environment.",
        style="Normal")
    doc.add_paragraph(style="Normal")

    # ---- Quick Overview (Heading 1) ----------------------------------------
    doc.add_heading("Quick Overview", level=1)

    _add_table(doc,
        headers=["TRL", "Name", "Core question answered", "Environment"],
        rows=[
            ["1", "Basic principles observed",
             '"Is this physically possible?"', "Pure research / paper"],
            ["2", "Technology concept formulated",
             '"What could we build with it?"', "Concept / paper"],
            ["3", "Experimental proof of concept",
             '"Does the key idea actually work?"', "Lab (analytical / small test)"],
            ["4", "Technology validated in lab",
             '"Do the parts work together?"', "Lab (breadboard)"],
            ["5", "Technology validated in relevant environment",
             '"Does it work in realistic conditions?"', "Simulated / relevant environment"],
            ["6", "Technology demonstrated in relevant environment",
             '"Does a full prototype work?"', "Relevant environment"],
            ["7", "System prototype in operational environment",
             '"Does it work in the real setting?"', "Operational environment"],
            ["8", "System complete and qualified",
             '"Is it finished, tested and certified?"', "Operational (final form)"],
            ["9", "Actual system proven in operations",
             '"Is it in real use and reliable?"', "Live operation / production"],
        ]
    )

    _add_para(doc, "A useful mental split:", style="Normal")
    _add_bullet(doc, "TRL 1–3 → Research (prove the idea).")
    _add_bullet(doc, "TRL 4–6 → Development (build and validate prototypes).")
    _add_bullet(doc, "TRL 7–9 → Deployment (qualify, release, operate).")
    doc.add_paragraph(style="Normal")

    # ---- Per-TRL sections --------------------------------------------------
    trl_sections = [
        {
            "title": "TRL 1 — Basic Principles Observed",
            "definition":
                "TRL 1 is the lowest level of technology maturity. Scientific research is just "
                "beginning to be translated into applied research and development: basic scientific "
                "principles are observed and reported, but no practical application has been "
                "identified yet. In the NASA and EU Horizon 2020 definitions this is the \"paper "
                "study\" level — you understand *that* an effect exists, not yet *how* to use it.",
            "goal":
                "Establish and document that the underlying scientific/engineering principle the "
                "future product will rely on genuinely exists and is understood.",
            "activities": [
                "Literature review and study of existing published research, standards and prior art.",
                "Observation and reporting of the basic physical, electrical, chemical or thermal "
                "principles that could be exploited.",
                "Purely theoretical or analytical work; any experiments only observe a phenomenon, "
                "they do not build anything.",
                "No application, requirement or design is defined yet — the work is knowledge-gathering.",
            ],
            "environment":
                "Entirely on paper / in the scientific literature. There is no hardware, no device "
                "model, and no defined operating environment.",
            "evidence":
                "Research notes or a short survey report that identifies the relevant principle(s), "
                "with references to papers, textbooks or standards. Reviewers should be able to "
                "answer *\"is this physically possible?\"* with a documented \"yes\".",
            "pitfalls":
                "Jumping to a solution before the principle is understood; confusing a known "
                "principle with a proven application (that is TRL 3+, not TRL 1).",
            "repo":
                "For the power-electronics projects here (400 A AC/DC charger, e-PU Cabinet V2 "
                "BESS) the basic principles — galvanic isolation, PFC rectification, DC/DC "
                "conversion, liquid cooling, Li-ion storage — are long-established and well "
                "documented. TRL 1 is therefore effectively \"prior art\"; it is assumed complete "
                "before the project starts and is not called out as an explicit task.",
        },
        {
            "title": "TRL 2 — Technology Concept Formulated",
            "definition":
                "Once basic principles are observed, practical applications can be invented. TRL 2 "
                "is where a *specific* technology concept and its intended application are "
                "formulated. The concept is still speculative — there is little or no experimental "
                "proof yet — but the intended use, the analytical reasoning and the expected "
                "benefit are written down.",
            "goal":
                "Turn the known principle into a defined, *possible* product concept: state what "
                "would be built, for whom, and against which requirements — even though nothing has "
                "been proven.",
            "activities": [
                "Define the concept and its intended use (the product idea and its application).",
                "Formulate the project scope, key results, and technical specification, split into "
                "must-have and nice-to-have requirements.",
                "Perform an initial risk analysis and a scope-change analysis.",
                "Carry out market / component research: which suppliers and devices exist, at what "
                "price, and with what documentation quality.",
            ],
            "environment":
                "Still analytical / on paper. The work consists of specifications, studies and "
                "comparisons rather than hardware.",
            "evidence":
                "A project/scope document, a written specification (requirements) list, and a "
                "market/component study. These let a reviewer answer *\"what could we build, and "
                "with which parts?\"*.",
            "pitfalls":
                "Writing requirements that are not measurable; skipping the must-have / nice-to-have "
                "split; underestimating supplier lead-time or documentation-quality risk.",
            "repo":
                "TRL 2 is the first explicit task in the project scope. For the 400 A charger it is "
                "where the must-have specification is fixed (galvanic isolation AC↔DC, 400/500 A "
                "Powerlock 3P+PE, 600–800 VDC, liquid cooled, material cost < €20 000, Modbus TCP, "
                "IEC 61439 / NEN 1010 / IEC 61000 compliance, C80 inrush) and where the module "
                "market study compares candidate converter modules and suppliers.",
        },
        {
            "title": "TRL 3 — Experimental Proof of Concept",
            "definition":
                "TRL 3 is where active research and development begins. Analytical studies and/or "
                "small laboratory experiments are used to physically validate that the *critical "
                "function* or analytical prediction of the concept actually works. It proves the "
                "single riskiest element — not the whole system.",
            "goal":
                "Prove that the *critical function* of the idea really works, analytically or with "
                "a small-scale test, and translate the concept into a first conceptual design.",
            "activities": [
                "Analytical studies and/or small-scale lab experiments focused on the key risky element.",
                "Conceptual design: decide which components are required and show how they satisfy "
                "the scope.",
                "First engineering artefacts: single-line electrical diagram, mechanical concept "
                "description.",
                "Material cost estimate and a rough engineering/planning (hours, cost, schedule) "
                "estimate based on the concept.",
            ],
            "environment":
                "Laboratory or analytical. Any hardware used is ad-hoc and not representative of "
                "the final product; the focus is on the proof, not on integration or packaging.",
            "evidence":
                "Proof-of-concept results (analysis or bench test), a conceptual design, a "
                "single-line diagram, and a cost/planning estimate. In this repo's flow the phase "
                "ends with a **tollgate review**.",
            "pitfalls":
                "Proving an easy part while ignoring the real risk; conceptual designs that ignore "
                "cost, size or cooling constraints; optimistic hour estimates.",
            "repo":
                "TRL 3 is the conceptual-design task: selecting converter modules, drawing the "
                "single-line diagram, describing the mechanical layout inside the power-module "
                "envelope, and estimating material cost (target < €20 000) and engineering hours. "
                "The tollgate decides whether the concept is worth detailed engineering.",
        },
        {
            "title": "TRL 4 — Technology Validated in the Laboratory",
            "definition":
                "With the concept proven, the basic technological components are integrated to "
                "establish that they will work together. At TRL 4 this integration is \"low "
                "fidelity\" — a laboratory breadboard — compared with the eventual system. It is "
                "the first level where the pieces are combined rather than tested in isolation.",
            "goal":
                "Show that the basic components function *together* (integration) in a controlled "
                "laboratory setting.",
            "activities": [
                "Build and test a simple prototype (\"breadboard\") that integrates the key components.",
                "Detailed design work begins in parallel (mechanical, electrical, software).",
                "Laboratory measurements confirm the combined parts behave as predicted.",
            ],
            "environment":
                "Controlled laboratory. The breadboard is functional but not packaged, ruggedised "
                "or representative of the final form factor.",
            "evidence":
                "Breadboard test results, the first detailed design drawings/schematics, and "
                "laboratory measurement reports.",
            "pitfalls":
                "Treating a tidy bench demo as if it were a relevant-environment test; deferring "
                "EMC, thermal and isolation questions that will dominate later levels.",
            "repo":
                "TRL 4 marks the start of detailed design. In this repo TRL 4 and 5 are handled "
                "together as one \"detailed design\" step, so TRL 4 covers the first integrated lab "
                "validation of the charger's power path (rectifier/PFC + isolated DC/DC + control) "
                "before it is exercised under realistic conditions.",
        },
        {
            "title": "TRL 5 — Technology Validated in a Relevant Environment",
            "definition":
                "The fidelity of the breadboard increases significantly. At TRL 5 the integrated "
                "technology is tested in a *relevant* (realistic, though possibly simulated) "
                "environment, so that it is validated under conditions much closer to the real "
                "application — for key enabling technologies, an industrially relevant environment.",
            "goal":
                "Increase fidelity — test the integrated technology under conditions that "
                "*resemble* real use (realistic loads, temperatures, vibration, EMC), not an ideal "
                "bench.",
            "activities": [
                "Test the prototype/subsystem in a simulated or relevant environment (realistic "
                "load, temperature, vibration, EMC conditions rather than an ideal bench).",
                "Achieve more complete integration than at TRL 4.",
                "Complete the detailed engineering: galvanic isolation, EMC, inrush, cooling, "
                "envelope, weight.",
                "Produce CE / compliance documentation, the build book and assembly instructions.",
            ],
            "environment":
                "Relevant environment — realistic operating stresses applied, though possibly still "
                "on a test rig rather than the final installation.",
            "evidence":
                "Validation results obtained in a relevant environment, a complete detailed design "
                "package, and draft CE documentation. (In this repo TRL 4 and 5 are handled "
                "together as \"detailed design\".) Ends with a **tollgate review**.",
            "pitfalls":
                "Declaring success from a single nominal-condition test; leaving compliance (CE, "
                "IEC 61439, EMC IEC 61000) documentation until after the hardware is frozen.",
            "repo":
                "TRL 5 completes detailed design for the charger: liquid-cooling sizing (total heat "
                "< 7 kW, efficiency > 97.5 %), galvanic isolation, C-type inrush behaviour, weight "
                "(< 500 kg) and envelope (within the power-module size), plus the CE dossier and "
                "build book. It feeds directly into the prototype build.",
        },
        {
            "title": "TRL 6 — Technology Demonstrated in a Relevant Environment",
            "definition":
                "TRL 6 is a major step up: a representative model or full system/subsystem "
                "prototype is built and demonstrated in a relevant environment. The prototype is "
                "well beyond the breadboard of TRL 5 and is close to the final configuration in "
                "form, fit and function.",
            "goal":
                "Demonstrate a full system/subsystem prototype in a relevant environment, validated "
                "against the must-have criteria.",
            "activities": [
                "Build the prototype (proto build).",
                "Functional testing and performance validation against the must-have criteria.",
                "Verify the key behaviours (EMC, inrush current, thermal behaviour, ingress "
                "protection, efficiency, control interface).",
            ],
            "environment":
                "Relevant environment with a near-representative prototype — much higher fidelity "
                "than TRL 5, exercising the real power levels and interfaces.",
            "evidence":
                "A working prototype, a functional test report, and a performance validation "
                "against requirements — exactly what a *measurement / verification plan* checks "
                "off. Ends with a **tollgate review**.",
            "pitfalls":
                "Testing only a subset of requirements; no traceable pass/fail matrix; skipping "
                "worst-case (maximum power, maximum temperature) points.",
            "repo":
                "TRL 6 is the testing task. The e-PU Cabinet V2 Measurement Plan is the TRL 6 "
                "instrument: it lists every requirement and records that the prototype was measured "
                "and passes (power, efficiency, isolation, EMC, inrush, cooling, Modbus TCP "
                "control, etc.).",
        },
        {
            "title": "TRL 7 — System Prototype Demonstration in an Operational Environment",
            "definition":
                "TRL 7 requires demonstration of a near-final (\"system prototype\") in the "
                "*operational* environment — the real setting in which the product will be used. It "
                "differs from TRL 6 in that the environment is the actual operational one, not a "
                "relevant/simulated one.",
            "goal":
                "Demonstrate a near-final prototype in the actual operational environment.",
            "activities": [
                "Field trials / pilot installation in the real operating setting (installed on "
                "site, connected to the real grid or microgrid).",
                "Full-scale operational testing under real-world conditions and duty cycles.",
            ],
            "environment":
                "Operational environment — the real installation, real grid/microgrid, real "
                "thermal and duty conditions. Fidelity is essentially final.",
            "evidence":
                "Field/pilot test results and an operational demonstration report from the real site.",
            "pitfalls":
                "Assuming a lab-passed prototype will behave identically on a real microgrid "
                "(earthing, harmonics, ambient); insufficient monitoring during the pilot.",
            "repo":
                "The projects here often move directly from TRL 6 (prototype test) to TRL 8 (sales "
                "readiness), so TRL 7 is not always an explicit task. It is where an on-site "
                "operational pilot of the charger in an e-PU10 microgrid would sit — validating the "
                "> 98 % uptime and serviceability goals under real conditions — if such a pilot is "
                "run.",
        },
        {
            "title": "TRL 8 — Actual System Completed and Qualified",
            "definition":
                "The technology has been proven to work in its final form and under the expected "
                "conditions. At TRL 8 the actual system is completed and *qualified* through test "
                "and demonstration; in almost all cases this is the end of true system development.",
            "goal":
                "Reach the final form — fully tested, qualified and certified, and ready to be "
                "offered commercially.",
            "activities": [
                "Final product and process qualification.",
                "Certification / compliance sign-off (e.g. CE marking).",
                "Full integration and pre-commercial preparation.",
                "Commercial preparation: sales one-pager, technical datasheet.",
            ],
            "environment":
                "Final form in the operational configuration; the product is the real thing, "
                "qualified against its specification and standards.",
            "evidence":
                "Qualification/certification records, a technical datasheet and a sales one-pager. "
                "Ends with a **tollgate review**.",
            "pitfalls":
                "Treating certification as a formality; datasheet figures that don't match the "
                "qualified test results; missing traceability from requirement → test → certificate.",
            "repo":
                "TRL 8 is the \"create one-pager for sales\" task. For the charger it means the "
                "CE-marked, qualified 400 A unit with a datasheet reflecting the verified "
                "specification (power, efficiency, isolation, cost, interfaces) ready to present to "
                "customers.",
        },
        {
            "title": "TRL 9 — Actual System Proven Through Successful Operations",
            "definition":
                "TRL 9 is the highest level: the actual system is proven through successful "
                "operation. The technology is applied in its final form and under real, routine "
                "operating conditions, and is no longer under development — only in-service "
                "monitoring and improvement remain.",
            "goal":
                "The technology is in real, routine operational use and proven reliable, and is "
                "released as a series product.",
            "activities": [
                "Finalize and update all documentation.",
                "Make the product series-ready (production-ready).",
                "Transfer the product to the sales matrix / into production.",
                "Ongoing monitoring and maintenance.",
            ],
            "environment":
                "Live operation / series production — the real product, in real use, at production "
                "quality.",
            "evidence":
                "A released, series-ready product; final documentation; the product in the sales "
                "catalogue; and field reliability data. Ends with a final acceptance **tollgate "
                "review**.",
            "pitfalls":
                "Declaring TRL 9 from a single successful install rather than proven routine "
                "operation; letting documentation drift out of date once the product ships; no "
                "feedback loop from field data into maintenance.",
            "repo":
                "TRL 9 is the \"finalize product and transfer to sales matrix\" task. For the 400 A "
                "charger and the e-PU Cabinet V2 BESS it means the production-ready unit, complete "
                "documentation, and the product in VDL Energy Systems' sales matrix, with field "
                "uptime (> 98 %) monitored in service.",
        },
    ]

    for section in trl_sections:
        doc.add_page_break()
        doc.add_heading(section["title"], level=1)
        _add_para(doc, section["definition"], style="Normal")
        doc.add_heading("Goal", level=2)
        _add_para(doc, section["goal"], style="Normal")
        doc.add_heading("What is normally done", level=2)
        for item in section["activities"]:
            _add_bullet(doc, item)
        doc.add_heading("Environment & fidelity", level=2)
        _add_para(doc, section["environment"], style="Normal")
        doc.add_heading("Typical output / evidence (exit criteria)", level=2)
        _add_para(doc, section["evidence"], style="Normal")
        doc.add_heading("Common pitfalls", level=2)
        _add_para(doc, section["pitfalls"], style="Normal")
        doc.add_heading("How this applies to this repository", level=2)
        _add_para(doc, section["repo"], style="Normal")
        doc.add_paragraph(style="Normal")

    # ---- Mapping to this repo (Heading 1) ----------------------------------
    doc.add_heading("How This Maps to the Projects in This Repo", level=1)
    _add_para(doc,
        "The engineering projects here use a slightly condensed TRL flow with tollgate reviews "
        "between phases:",
        style="Normal")
    _add_bullet(doc, "TRL 2 – Scope & market research (scope, spec, risk analysis, component market study).")
    _add_bullet(doc, "TRL 3 – Conceptual design (component selection, single-line diagram, mechanical concept, cost/planning).")
    _add_bullet(doc, "TRL 4/5 – Detailed design (mechanical, electrical, software, CE documentation, build book).")
    _add_bullet(doc, "TRL 6 – Prototype build & test (proto build, functional/performance/EMC/thermal validation).")
    _add_bullet(doc, "TRL 8 – Sales readiness (one-pager, technical datasheet).")
    _add_bullet(doc, "TRL 9 – Product release (finalize docs, make series-ready, transfer to sales matrix).")
    doc.add_paragraph(style="Normal")
    _add_para(doc,
        "(TRL 1 and TRL 7 are usually not called out explicitly: TRL 1 is prior research, and the "
        "on-site operational pilot of TRL 7 is effectively folded into the prototype test and "
        "product release steps.)",
        style="Normal")
    doc.add_paragraph(style="Normal")

    # ---- References (Heading 1) --------------------------------------------
    doc.add_heading("References", level=1)
    _add_bullet(doc,
        "NASA Technology Readiness Level definitions — "
        "https://www.nasa.gov/directorates/heo/scan/engineering/technology/technology_readiness_level")
    _add_bullet(doc,
        "European Commission / Horizon 2020 TRL definitions (Annex G) — "
        "https://ec.europa.eu/research/participants/data/ref/h2020/wp/2014_2015/annexes/"
        "h2020-wp1415-annex-g-trl_en.pdf")

    # ---- Save --------------------------------------------------------------
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    print(f"Saved: {output_path}")


if __name__ == "__main__":
    build(TEMPLATE, SOURCE_MD, OUTPUT)
