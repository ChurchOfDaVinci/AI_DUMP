#!/usr/bin/env python3
"""
generate_measurement_plan.py
============================
Generates 'ACDC charger/400A Charger Measurement Plan.docx' by:
  1. Copying the template (styles, headers/footers, theme, media, numbering) from
     'ACDC charger/e-PU Cabinet V2 Measurement Plan.docx'
  2. Clearing the document body while preserving all style / layout artefacts
  3. Writing new content that mirrors the e-PU Cabinet V2 measurement plan chapter
     structure (Introduction + How to use, Legend, Verification Cross-Reference
     Matrix, Test Sheets grouped per TRL, Test Phase Overview, Sign-off) for the
     400A AC/DC charger project.

Usage
-----
  cd <repo root>
  python3 scripts/generate_measurement_plan.py

Dependencies
------------
  pip install python-docx
"""

import copy
import io
import os
import shutil
import zipfile
from datetime import date

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
REPO_ROOT = os.path.join(os.path.dirname(__file__), "..")
TEMPLATE_PATH = os.path.join(REPO_ROOT, "ACDC charger", "e-PU Cabinet V2 Measurement Plan.docx")
OUTPUT_PATH = os.path.join(REPO_ROOT, "ACDC charger", "400A Charger Measurement Plan.docx")

# ---------------------------------------------------------------------------
# Helper – low-level XML utilities
# ---------------------------------------------------------------------------

def _set_cell_bg(cell, hex_color: str):
    """Set table-cell background shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    existing = tcPr.find(qn("w:shd"))
    if existing is not None:
        tcPr.remove(existing)
    tcPr.append(shd)


def _bold_cell(cell):
    """Make all runs in a cell bold."""
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
        if not para.runs:
            run = para.add_run(para.text)
            run.bold = True


def _make_row_header(row, hex_color: str = "1F3864"):
    """Apply dark-blue background + white bold text to a table row."""
    for cell in row.cells:
        _set_cell_bg(cell, hex_color)
        for para in cell.paragraphs:
            for run in para.runs:
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            if not para.runs and para.text:
                run = para.add_run(para.text)
                run.bold = True
                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def _make_label_cell(cell, hex_color: str = "D9E1F2"):
    """Light-blue background for label cells."""
    _set_cell_bg(cell, hex_color)
    for para in cell.paragraphs:
        for run in para.runs:
            run.bold = True
        if not para.runs and para.text:
            run = para.add_run(para.text)
            run.bold = True


def _add_paragraph(doc, text: str, style: str = "Normal", bold: bool = False) -> object:
    para = doc.add_paragraph(style=style)
    if text:
        run = para.add_run(text)
        if bold:
            run.bold = True
    return para


def _page_break(doc):
    para = doc.add_paragraph()
    run = para.add_run()
    run.add_break(docx_break_type=1)  # WD_BREAK.PAGE
    return para


# ---------------------------------------------------------------------------
# Test-sheet data model
# ---------------------------------------------------------------------------

class TestSheet:
    def __init__(self, ts_id, title, objective, standard, equipment,
                 setup, procedure_steps, measured_qty, acceptance, method="T"):
        self.ts_id = ts_id
        self.title = title
        self.objective = objective
        self.standard = standard
        self.equipment = equipment
        self.setup = setup
        self.procedure_steps = procedure_steps  # list of strings
        self.measured_qty = measured_qty
        self.acceptance = acceptance
        self.method = method


# ---------------------------------------------------------------------------
# All test-sheet content (derived from project.txt, TRL_Tasks and SCORE csv)
# ---------------------------------------------------------------------------

TRL_DATA = [
    # -----------------------------------------------------------------------
    # TRL 2 — Scope & Market Research
    # -----------------------------------------------------------------------
    {
        "number": "2",
        "title": "Scope & Market Research",
        "weeks": "Week 1–3",
        "effort": "32 h",
        "objective": (
            "Confirm the project scope, define and record all technical requirements "
            "(must-have and recommended), perform a component market survey to identify "
            "available AC/DC converter modules that can satisfy the 400 A charger specification, "
            "and assess the quality of supplier documentation."
        ),
        "entry_criteria": [
            "Project kick-off meeting has taken place (23-06-2026).",
            "Project plan document (ProjectPlan_400A_Charger.docx) is available and baselined.",
        ],
        "exit_criteria": [
            "Project-scope document is finalised and approved by the project lead.",
            "Technical requirements list (must-have / recommended) is complete.",
            "At least three candidate converter modules have been evaluated against the requirements.",
            "Risk register has been created.",
            "Tollgate TRL 2 review is signed off by project lead (L. Rietkerk).",
        ],
        "sheets": [
            TestSheet(
                ts_id="TS-2-01",
                title="Must-Have Specification Compliance Checklist",
                objective=(
                    "Verify that the project scope document contains all mandatory requirements "
                    "and that each requirement is traceable to the project plan."
                ),
                standard="Internal project plan; IEC 61439; NEN 1010",
                equipment="Project plan document; requirements checklist template",
                setup=(
                    "Review the project plan document against the must-have list. "
                    "Perform at desk; no physical equipment required."
                ),
                procedure_steps=[
                    "Open ProjectPlan_400A_Charger.docx.",
                    "Check that each must-have item in the list below has a corresponding entry in the project plan:",
                    "  a. Galvanic AC/DC isolation",
                    "  b. 400/500 A Powerlock (3P+PE)",
                    "  c. Unidirectional operation",
                    "  d. Liquid cooling",
                    "  e. DC voltage range ≥ 600–800 VDC",
                    "  f. Material cost < €20,000",
                    "  g. Communication: ModbusTCP / CAN / Profinet / EtherCAT",
                    "  h. EU compliant",
                    "  i. Mobile / vibration-resistant",
                    "  j. Min. settable power ≥ 55 kW / 80 A",
                    "  k. Inrush acceptable for C-type C80 breaker",
                    "Tick each item as Present / Missing in the result column.",
                    "Record any missing items as non-conformances.",
                ],
                measured_qty="Number of must-have requirements addressed (count / 11)",
                acceptance="All 11 must-have requirements are present and unambiguously formulated. No missing items.",
                method="I / A",
            ),
            TestSheet(
                ts_id="TS-2-02",
                title="Supplier & Component Documentation Quality Assessment",
                objective=(
                    "Evaluate the availability and completeness of technical documentation "
                    "for each candidate AC/DC converter module identified in the market survey."
                ),
                standard="Internal quality criteria; IEC 61439",
                equipment="Supplier datasheets, test reports, certificates; SCORE evaluation matrix (400A - 275kW AC-DC Converters(SCORE).csv)",
                setup=(
                    "Collect datasheets and certificate documents for all candidate modules. "
                    "Use the SCORE spreadsheet to rate each supplier."
                ),
                procedure_steps=[
                    "List all candidate modules identified in the market survey.",
                    "For each module, check the availability of:",
                    "  a. Complete datasheet (power, voltage range, efficiency, dimensions, weight)",
                    "  b. CE / TUV certificate or declaration of conformity",
                    "  c. Communication protocol documentation (register map / CAN DBC)",
                    "  d. Mechanical drawing or 3D model",
                    "  e. Cooling specifications (flow rate, pressure drop, inlet temperature)",
                    "Score each criterion 0 (missing) / 5 (partial) / 10 (complete).",
                    "Calculate overall documentation score per candidate.",
                    "Record results in SCORE matrix.",
                ],
                measured_qty="Documentation completeness score per candidate (0–10 per criterion)",
                acceptance=(
                    "At least two candidates score ≥ 7.0 overall. "
                    "Each selected candidate has a valid CE/TUV declaration on file."
                ),
                method="I / A",
            ),
            TestSheet(
                ts_id="TS-2-03",
                title="Component Availability vs. Requirements Verification",
                objective=(
                    "Confirm that at least one commercially available converter module meets "
                    "all must-have requirements from the 400 A charger specification."
                ),
                standard="Internal project specification; IEC 61439; NEN 1010",
                equipment="SCORE evaluation matrix; supplier datasheets",
                setup="Cross-reference SCORE evaluation results against the must-have checklist (TS-2-01).",
                procedure_steps=[
                    "Take the ranked list of candidates from TS-2-02.",
                    "For each candidate, verify compliance with each must-have requirement.",
                    "Mark each requirement Pass (P) / Fail (F) / Unknown (U) per candidate.",
                    "Identify the top-scoring candidate(s) with ≤ 1 Unknown item.",
                    "Document any gaps or assumptions for further investigation in TRL 3.",
                ],
                measured_qty="Number of must-have requirements met per candidate (count / 11)",
                acceptance=(
                    "At least one candidate meets all 11 must-have requirements or has at most "
                    "one Unknown item (to be resolved in TRL 3)."
                ),
                method="A",
            ),
        ],
    },
    # -----------------------------------------------------------------------
    # TRL 3 — Conceptual Design
    # -----------------------------------------------------------------------
    {
        "number": "3",
        "title": "Conceptual Design",
        "weeks": "Week 4–5",
        "effort": "20 h",
        "objective": (
            "Select the required components, define how the conceptual design meets the "
            "project scope, create a single-line electrical diagram and mechanical concept, "
            "and validate the design concept against cost, thermal and mechanical constraints."
        ),
        "entry_criteria": [
            "TRL 2 tollgate is signed off.",
            "At least one candidate module identified and short-listed.",
            "Technical requirements list is complete.",
        ],
        "exit_criteria": [
            "Single-line electrical diagram is reviewed and approved.",
            "Mechanical concept description is complete (dimensions, weight estimate).",
            "Material cost estimate is within budget.",
            "Thermal/efficiency budget confirms the design is feasible.",
            "Tollgate TRL 3 review is signed off by project lead.",
        ],
        "sheets": [
            TestSheet(
                ts_id="TS-3-01",
                title="Single-Line Electrical Design Review",
                objective=(
                    "Verify that the single-line electrical diagram correctly represents the "
                    "400 A charger architecture and includes all required functional blocks."
                ),
                standard="IEC 61439; NEN 1010; internal design standard",
                equipment="Single-line diagram (CAD/PDF); design review checklist",
                setup="Print or display the draft single-line diagram for a structured design review.",
                procedure_steps=[
                    "Obtain the draft single-line electrical diagram.",
                    "Verify the following elements are present and correctly connected:",
                    "  a. AC input: 3P+PE, 400/500 A Powerlock connector",
                    "  b. AC input protection: circuit breaker rated for C80 (C-type)",
                    "  c. Galvanic isolation barrier (transformer or equivalent)",
                    "  d. AC/DC converter module(s)",
                    "  e. DC output bus: 600–800 VDC range",
                    "  f. DC output protection (fuse / switch)",
                    "  g. Liquid-cooling connections",
                    "  h. Control / communication interface (ModbusTCP / CAN)",
                    "  i. Earth bonding and protective-earth routing",
                    "Record any missing elements or design errors.",
                    "Sign the diagram as reviewed.",
                ],
                measured_qty="Number of required elements present and correctly connected (count / 9)",
                acceptance="All 9 required elements are present. Zero design errors recorded. Diagram is signed off.",
                method="I / A",
            ),
            TestSheet(
                ts_id="TS-3-02",
                title="Thermal & Efficiency Budget Check",
                objective=(
                    "Verify by calculation that the conceptual design achieves the required "
                    "efficiency and stays within the e-PU10 heat dissipation limit."
                ),
                standard="e-PU10 restrictions (project plan); IEC 61439",
                equipment="Calculator / spreadsheet; converter efficiency data from datasheets",
                setup=(
                    "Gather efficiency figures from the selected converter module datasheet. "
                    "Determine system-level losses (converter, cabling, connectors)."
                ),
                procedure_steps=[
                    "Obtain full-load efficiency η_converter from datasheet (at rated output power).",
                    "Estimate cabling + connector losses: P_cable [W].",
                    "Calculate total input power: P_in = P_out / (η_converter) + P_cable.",
                    "Calculate total heat dissipation: P_heat = P_in − P_out [W].",
                    "Calculate system efficiency: η_sys = P_out / P_in × 100 [%].",
                    "Record η_sys and P_heat in the result field.",
                    "Compare against limits: η_sys ≥ 97.5 % and P_heat < 7 000 W.",
                ],
                measured_qty="System efficiency η_sys [%]; total heat P_heat [W]",
                acceptance="η_sys ≥ 97.5 % AND P_heat < 7 000 W at rated output power.",
                method="A",
            ),
            TestSheet(
                ts_id="TS-3-03",
                title="Mechanical Envelope & Weight Pre-Check",
                objective=(
                    "Verify that the conceptual mechanical design fits within the e-PU10 "
                    "power-module envelope and does not exceed the maximum weight."
                ),
                standard="e-PU10 restrictions (project plan); mechanical concept drawing",
                equipment="Mechanical concept sketch / CAD model; tape measure (if physical mock-up exists)",
                setup="Use the mechanical concept description and component datasheets to estimate total dimensions and weight.",
                procedure_steps=[
                    "Obtain mechanical dimensions of the selected converter module(s) from datasheet.",
                    "Add housing / enclosure dimensions (estimated).",
                    "Check that overall assembly fits within power-module size envelope (e-PU10 constraint).",
                    "Sum weights of all components (modules + enclosure + cooling system + cabling).",
                    "Record total estimated weight [kg].",
                    "Compare against limit: weight < 500 kg.",
                ],
                measured_qty="Overall assembly dimensions [mm × mm × mm]; estimated total weight [kg]",
                acceptance=(
                    "Assembly dimensions fit within e-PU10 power-module size. "
                    "Estimated total weight < 500 kg."
                ),
                method="A / I",
            ),
            TestSheet(
                ts_id="TS-3-04",
                title="Material Cost Check",
                objective=(
                    "Verify that the estimated bill-of-materials (BOM) cost of the conceptual "
                    "design is within budget."
                ),
                standard="Project plan (cost constraint < €20,000)",
                equipment="Component quotations / list prices; BOM spreadsheet",
                setup="Compile a preliminary BOM from the conceptual design and obtain unit prices from TS-2-02 supplier list.",
                procedure_steps=[
                    "List all major components: converter module(s), enclosure, cooling components, connectors, cables, protection devices, control hardware.",
                    "Obtain unit price for each component (quotation or list price).",
                    "Calculate total BOM cost: Σ (quantity × unit price) [€].",
                    "Add 15 % contingency for unknown items.",
                    "Record total estimated material cost [€].",
                    "Compare against limit: < €20,000.",
                ],
                measured_qty="Total estimated material cost [€]",
                acceptance="Total estimated BOM cost (including 15 % contingency) < €20,000.",
                method="A",
            ),
        ],
    },
    # -----------------------------------------------------------------------
    # TRL 4/5 — Detailed Design
    # -----------------------------------------------------------------------
    {
        "number": "4.5",
        "title": "Detailed Design",
        "weeks": "Week 6–12",
        "effort": "180 h",
        "objective": (
            "Execute the full detailed design: mechanical, electrical and software. "
            "Verify on paper / by analysis that the detailed design satisfies all must-have "
            "requirements before prototype build is authorised."
        ),
        "entry_criteria": [
            "TRL 3 tollgate is signed off.",
            "Conceptual design is approved.",
            "Component selection is confirmed (at least one module meets all must-haves).",
        ],
        "exit_criteria": [
            "All six detailed-design test sheets are passed (or accepted with derogation).",
            "Detailed mechanical and electrical drawings are released.",
            "Software architecture and communication register map are documented.",
            "CE documentation package is started.",
            "Tollgate TRL 4/5 review is signed off by project lead.",
        ],
        "sheets": [
            TestSheet(
                ts_id="TS-4.5-01",
                title="Galvanic Isolation / Dielectric Withstand Verification",
                objective=(
                    "Verify by design analysis that the AC/DC converter module provides "
                    "galvanic isolation between AC input and DC output, and that the isolation "
                    "level meets IEC 61439 / NEN 1010 requirements."
                ),
                standard="IEC 61439-1; NEN 1010; converter module datasheet / test report",
                equipment="Converter datasheet; IEC 61439 standard (design analysis – no physical test required at this TRL)",
                setup=(
                    "Review the converter module's isolation test report or certificate. "
                    "Confirm isolation voltage, creepage and clearance from the datasheet."
                ),
                procedure_steps=[
                    "Obtain the converter module's dielectric withstand / hi-pot test report.",
                    "Check that the reported test voltage meets or exceeds the IEC 61439 requirement for the rated working voltage.",
                    "Verify creepage and clearance values against IEC 61439 Table F.2.",
                    "Confirm that the module is described as 'galvanically isolated' in the datasheet.",
                    "Record: test voltage [V], isolation class, creepage [mm], clearance [mm].",
                    "If no test report is available, flag as action item for TRL 6 physical test.",
                ],
                measured_qty="Dielectric withstand voltage [V AC or V DC]; isolation class; creepage [mm]; clearance [mm]",
                acceptance=(
                    "Dielectric withstand ≥ 3 750 V AC (or per IEC 61439 for rated voltage). "
                    "Galvanic isolation confirmed in datasheet. "
                    "Creepage and clearance within IEC 61439 limits."
                ),
                method="I / A",
            ),
            TestSheet(
                ts_id="TS-4.5-02",
                title="DC Output Voltage Range Verification",
                objective=(
                    "Verify by design analysis that the converter module can operate across "
                    "the required DC output voltage range of 600–800 VDC (minimum)."
                ),
                standard="Project specification; converter module datasheet",
                equipment="Converter datasheet; voltage range specification from project plan",
                setup="Review datasheet output voltage range specification.",
                procedure_steps=[
                    "Obtain the DC output voltage range from the converter module datasheet.",
                    "Record minimum and maximum DC output voltage [V].",
                    "Verify minimum ≤ 600 V and maximum ≥ 800 V.",
                    "Check that the voltage regulation accuracy meets requirements (e.g. ± 1 %).",
                    "Confirm that the range is achievable via the communication interface (register / command).",
                ],
                measured_qty="DC output voltage range: V_min [V DC] and V_max [V DC]",
                acceptance="V_min ≤ 600 V DC AND V_max ≥ 800 V DC as specified in datasheet.",
                method="I / A",
            ),
            TestSheet(
                ts_id="TS-4.5-03",
                title="Minimum Settable Power Setpoint Verification",
                objective=(
                    "Verify by design analysis that the converter module supports a minimum "
                    "settable power setpoint of at least 55 kW / 80 A (must-have)."
                ),
                standard="Project specification (must-have); converter module communication datasheet",
                equipment="Converter communication register map / datasheet",
                setup="Review the communication register map for minimum power / current setpoint resolution.",
                procedure_steps=[
                    "Obtain the minimum power / current setpoint value from the datasheet or register map.",
                    "Record P_min [kW] and I_min [A].",
                    "Verify P_min ≤ 55 kW and I_min ≤ 80 A (i.e. the charger CAN be set as low as these values).",
                    "Check setpoint resolution (step size) is ≤ 1 kW or ≤ 1 A.",
                    "Additionally check if recommended minimum (20 kW / 32 A) is achievable.",
                ],
                measured_qty="Minimum settable power P_min [kW]; minimum settable current I_min [A]",
                acceptance=(
                    "P_min ≤ 55 kW AND I_min ≤ 80 A (must-have met). "
                    "P_min ≤ 20 kW AND I_min ≤ 32 A is preferred (recommended spec)."
                ),
                method="I / A",
            ),
            TestSheet(
                ts_id="TS-4.5-04",
                title="Communication Protocol Verification",
                objective=(
                    "Verify that the detailed software design implements a supported communication "
                    "protocol (ModbusTCP preferred; alternatively CAN, Profinet or EtherCAT) "
                    "and that a complete register map / message definition is documented."
                ),
                standard="Project specification; Modbus Application Protocol v1.1b3; CAN ISO 11898",
                equipment=(
                    "Software architecture document; register map / DBC file; "
                    "Modbus analyser (Modbus Poll or equivalent) — for TRL 6 physical test"
                ),
                setup=(
                    "Review the software architecture document and communication register map. "
                    "Confirm that the implemented protocol matches a supported option."
                ),
                procedure_steps=[
                    "Identify the communication protocol implemented in the detailed design.",
                    "Verify the protocol is one of: ModbusTCP (preferred), CAN, Profinet, EtherCAT.",
                    "Obtain the complete register map (Modbus) or DBC file (CAN).",
                    "Check that all required control commands are mapped:",
                    "  a. Enable / disable charging",
                    "  b. Set output power [kW] or current [A]",
                    "  c. Set output voltage [V]",
                    "  d. Read measured output power, current, voltage",
                    "  e. Read fault / status word",
                    "Verify that the register map is versioned and included in the documentation package.",
                ],
                measured_qty=(
                    "Protocol type (ModbusTCP / CAN / Profinet / EtherCAT); "
                    "number of required commands mapped (count / 5)"
                ),
                acceptance=(
                    "Protocol is one of the four accepted types. "
                    "All 5 required control commands are mapped and documented."
                ),
                method="I / A",
            ),
            TestSheet(
                ts_id="TS-4.5-05",
                title="EMC Pre-Compliance Design Review",
                objective=(
                    "Verify by design analysis that the detailed design incorporates appropriate "
                    "EMC mitigation measures to meet industrial EMC requirements (IEC 61000)."
                ),
                standard="IEC 61000-4 series; IEC 61000-6-2 (industrial immunity); IEC 61000-6-4 (industrial emission)",
                equipment="Detailed electrical drawings; EMC design checklist; converter EMC test report (if available)",
                setup=(
                    "Review detailed electrical drawings and the converter module's EMC certificate "
                    "(if available) against the IEC 61000 industrial EMC class requirements."
                ),
                procedure_steps=[
                    "Check that the converter module has an EMC test report or CE declaration covering IEC 61000.",
                    "Review the electrical drawings for the following EMC design features:",
                    "  a. Input EMC filter (line filter) is specified",
                    "  b. Cable shielding plan is documented",
                    "  c. PE / ground bonding topology is defined",
                    "  d. Separation of high-current power cables from signal cables is planned",
                    "Identify any gaps in EMC design and log as action items.",
                    "Record the EMC class of the converter module from its datasheet/certificate.",
                ],
                measured_qty="EMC class of converter (Class A / B); number of EMC design measures implemented (count / 4)",
                acceptance=(
                    "Converter module EMC class is industrial (IEC 61000-6-2/6-4 or equivalent). "
                    "All 4 EMC design measures are addressed in the design."
                ),
                method="I / A",
            ),
            TestSheet(
                ts_id="TS-4.5-06",
                title="Liquid-Cooling Circuit Design Verification",
                objective=(
                    "Verify by design analysis that the liquid-cooling circuit is correctly "
                    "designed to remove the required heat dissipation within the e-PU10 "
                    "cooling system constraints."
                ),
                standard="e-PU10 restrictions (project plan); converter module cooling datasheet",
                equipment=(
                    "Converter module cooling specification (flow rate, inlet temperature, pressure drop); "
                    "e-PU10 cooling system specification; hydraulic calculation spreadsheet"
                ),
                setup=(
                    "Obtain the converter module cooling requirements from the datasheet. "
                    "Obtain the e-PU10 cooling system parameters (available flow rate, max pressure drop, inlet temperature)."
                ),
                procedure_steps=[
                    "Record required coolant flow rate from converter datasheet [L/min].",
                    "Record maximum allowable pressure drop [bar].",
                    "Record maximum inlet coolant temperature [°C].",
                    "Verify that the e-PU10 cooling system can supply the required flow rate and inlet temperature.",
                    "Calculate heat removal capacity: Q = ṁ × c_p × ΔT [W], where ΔT = T_out − T_in.",
                    "Verify Q ≥ P_heat from TS-3-02 (< 7 000 W).",
                    "Check that all cooling connections (inlet, outlet, venting) are included in the mechanical drawing.",
                ],
                measured_qty="Coolant flow rate [L/min]; pressure drop [bar]; inlet temperature [°C]; heat removal capacity Q [W]",
                acceptance=(
                    "Cooling circuit supplies sufficient flow and temperature to remove ≥ 7 000 W. "
                    "Pressure drop within e-PU10 cooling system limits. "
                    "All cooling connections are shown in mechanical drawings."
                ),
                method="A / I",
            ),
        ],
    },
    # -----------------------------------------------------------------------
    # TRL 6 — Prototype Build & Test
    # -----------------------------------------------------------------------
    {
        "number": "6",
        "title": "Prototype Build & Test",
        "weeks": "Week 13–16",
        "effort": "56 h",
        "objective": (
            "Build the first physical prototype of the 400 A AC/DC charger and perform "
            "all physical validation tests to confirm that the prototype meets the must-have "
            "requirements. Results are recorded in this measurement plan."
        ),
        "entry_criteria": [
            "TRL 4/5 tollgate is signed off.",
            "All detailed drawings are released.",
            "Component procurement is complete.",
            "Test environment (AC supply, load bank, measurement instruments) is available and calibrated.",
        ],
        "exit_criteria": [
            "All eight TRL 6 test sheets are completed (Pass or accepted deviation with corrective action).",
            "Test results are recorded and signed off by the test engineer.",
            "Any deviations are documented with root cause and corrective action plan.",
            "Tollgate TRL 6 review is signed off by project lead.",
        ],
        "sheets": [
            TestSheet(
                ts_id="TS-6-01",
                title="Full-Load Continuous Operation at 25 °C",
                objective=(
                    "Verify that the charger delivers 100 % of rated output power continuously "
                    "at 25 °C ambient temperature without fault, shutdown or de-rating."
                ),
                standard="Project acceptance criteria; IEC 61439",
                equipment=(
                    "Calibrated AC power analyser (e.g. Yokogawa WT series or equivalent, calibrated ≤ 12 months); "
                    "calibrated DC power analyser; "
                    "resistive/electronic load bank rated ≥ rated charger output power; "
                    "temperature probe (calibrated); timer"
                ),
                setup=(
                    "Connect charger to a 3-phase 400 V AC supply via a C80 breaker. "
                    "Connect DC output to a calibrated electronic load. "
                    "Ambient temperature shall be 25 °C ± 3 °C (verify with calibrated thermometer). "
                    "Allow system to stabilise for 10 minutes before test."
                ),
                procedure_steps=[
                    "Set DC load to rated output voltage (within 600–800 VDC range, as specified).",
                    "Command charger to 100 % rated power via communication interface.",
                    "Start timer.",
                    "Record: AC input power [kW], DC output power [kW], DC output voltage [V], DC output current [A], ambient temperature [°C] at T = 0, 15, 30, 45, 60 min.",
                    "Monitor for any fault indication, shutdown or thermal de-rating event.",
                    "After 60 min continuous operation, record final values.",
                    "Stop charger and record any alarm codes.",
                ],
                measured_qty="DC output power [kW]; DC output current [A]; DC output voltage [V]; ambient temperature [°C]; duration without fault [min]",
                acceptance=(
                    "DC output power ≥ rated power (100 %) for ≥ 60 min without fault, "
                    "thermal de-rating or shutdown at 25 °C ± 3 °C ambient."
                ),
                method="T",
            ),
            TestSheet(
                ts_id="TS-6-02",
                title="Efficiency Measurement (Input vs. Output Power)",
                objective=(
                    "Measure the AC-to-DC conversion efficiency at 100 % rated load and "
                    "verify that the system efficiency exceeds 97.5 % (e-PU10 restriction)."
                ),
                standard="e-PU10 restrictions (project plan); IEC 61439",
                equipment=(
                    "Calibrated AC power analyser (≤ 12 months calibration, ±0.1 % accuracy); "
                    "calibrated DC power analyser (same accuracy); "
                    "electronic load bank"
                ),
                setup=(
                    "Same setup as TS-6-01. Both AC and DC power analysers shall be connected "
                    "simultaneously. Ensure all auxiliary power (fans, controls) is included in the AC measurement."
                ),
                procedure_steps=[
                    "With charger running at 100 % rated load (as per TS-6-01, steady state):",
                    "Record simultaneous AC input power P_AC_in [W] from AC analyser.",
                    "Record simultaneous DC output power P_DC_out [W] from DC analyser.",
                    "Calculate system efficiency: η = P_DC_out / P_AC_in × 100 [%].",
                    "Repeat measurement at 75 % and 50 % rated load.",
                    "Record efficiency at each load point.",
                    "Calculate total heat dissipation: P_heat = P_AC_in − P_DC_out [W] at 100 % load.",
                ],
                measured_qty="P_AC_in [W]; P_DC_out [W]; η [%]; P_heat [W]",
                acceptance=(
                    "η ≥ 97.5 % at 100 % rated load. "
                    "P_heat < 7 000 W at 100 % rated load."
                ),
                method="T",
            ),
            TestSheet(
                ts_id="TS-6-03",
                title="AC Inrush Behaviour under C80 / B32 Circuit Breaker",
                objective=(
                    "Verify that the inrush current on AC startup is acceptable for the "
                    "specified C-type C80 breaker (must-have) and preferably also for a "
                    "B-type B32 breaker (recommended)."
                ),
                standard="IEC 60947-2; project specification (must-have: C80; recommended: B32); IEC 60898-1",
                equipment=(
                    "Calibrated current clamp or Rogowski coil (bandwidth ≥ 10 kHz, calibrated); "
                    "oscilloscope (≥ 100 MHz, ≥ 4 channels); "
                    "C-type 80 A circuit breaker (for must-have test); "
                    "B-type 32 A circuit breaker (for recommended test)"
                ),
                setup=(
                    "Install the C80 breaker in the AC supply circuit. "
                    "Connect current clamps on all three phase conductors. "
                    "Charger DC output connected to electronic load set to rated voltage (open circuit mode initially). "
                    "Trigger oscilloscope on voltage rising edge of AC supply switch-on."
                ),
                procedure_steps=[
                    "With C80 breaker installed and AC supply disconnected from charger:",
                    "Switch on AC supply (close C80 breaker). Record inrush waveform on all three phases.",
                    "Measure peak inrush current I_peak [A] and duration t_inrush [ms].",
                    "Verify C80 breaker did not trip.",
                    "Repeat test 3 times with 5-minute cooling intervals. Record each result.",
                    "If B32 breaker is available: install B32 and repeat steps 1–4 for recommended spec.",
                    "Record all peak values and trip/no-trip outcome.",
                ],
                measured_qty="Peak inrush current I_peak [A]; inrush duration t [ms]; breaker trip status (trip / no trip)",
                acceptance=(
                    "MUST-HAVE: C80 breaker does not trip on any of 3 cold-start attempts. "
                    "RECOMMENDED: B32 breaker does not trip on any of 3 cold-start attempts."
                ),
                method="T",
            ),
            TestSheet(
                ts_id="TS-6-04",
                title="Thermal / Liquid-Cooling Performance at Full Load",
                objective=(
                    "Verify that the liquid-cooling circuit maintains component temperatures "
                    "within limits at 100 % rated load and that total heat dissipation is < 7 kW."
                ),
                standard="e-PU10 restrictions; converter module datasheet (max component temperature)",
                equipment=(
                    "Calibrated thermocouple logger (≥ 4 channels, ±0.5 °C, calibrated); "
                    "calibrated flow meter on cooling circuit; "
                    "calibrated inlet/outlet temperature sensors"
                ),
                setup=(
                    "Install thermocouples on: converter module heatsink, inlet coolant, "
                    "outlet coolant, and ambient. "
                    "Connect calibrated flow meter in the cooling loop."
                ),
                procedure_steps=[
                    "Start cooling pump and verify flow rate [L/min].",
                    "Run charger at 100 % rated load for 60 minutes (same run as TS-6-01).",
                    "Record every 10 minutes: heatsink temperature T_hs [°C], coolant inlet T_in [°C], coolant outlet T_out [°C], flow rate Q_cool [L/min], ambient T_amb [°C].",
                    "At steady state (temperature change < 1 °C over last 10 min), record all values.",
                    "Calculate heat removal: P_removed = ṁ × c_p × (T_out − T_in) [W], where ṁ = Q_cool × ρ_water.",
                    "Compare T_hs against converter module's maximum rated heatsink temperature.",
                ],
                measured_qty="T_hs [°C]; T_in [°C]; T_out [°C]; Q_cool [L/min]; P_removed [W]",
                acceptance=(
                    "T_hs ≤ converter module rated maximum heatsink temperature. "
                    "P_removed ≥ P_heat from TS-6-02 (i.e. cooling is sufficient). "
                    "No thermal alarm or shutdown during 60-minute full-load run."
                ),
                method="T",
            ),
            TestSheet(
                ts_id="TS-6-05",
                title="EMC Emissions and Immunity",
                objective=(
                    "Verify that the assembled charger prototype meets industrial EMC "
                    "requirements for emissions and immunity."
                ),
                standard=(
                    "IEC 61000-6-2 (industrial immunity); "
                    "IEC 61000-6-4 (industrial emission); "
                    "IEC 61000-4-2 (ESD); IEC 61000-4-4 (EFT); IEC 61000-4-5 (surge)"
                ),
                equipment=(
                    "EMC test laboratory (accredited preferred) or semi-anechoic test chamber; "
                    "spectrum analyser; LISN (line impedance stabilisation network); "
                    "ESD gun; EFT/burst generator; surge generator"
                ),
                setup=(
                    "Assemble prototype in standard operating configuration. "
                    "Connect to rated AC supply and rated DC load. "
                    "Perform tests in an accredited EMC lab or on-site pre-compliance."
                ),
                procedure_steps=[
                    "Conducted emissions test per IEC 61000-6-4 / CISPR 11 Class A.",
                    "Radiated emissions test per IEC 61000-6-4 Class A.",
                    "ESD immunity test per IEC 61000-4-2 (Level 3: 6 kV contact, 8 kV air).",
                    "Electrical fast transient / burst per IEC 61000-4-4 (Level 3).",
                    "Surge immunity per IEC 61000-4-5 (Level 3: 1 kV differential, 2 kV common mode).",
                    "Record pass/fail for each sub-test.",
                    "Log any failures with frequency, level and description.",
                ],
                measured_qty=(
                    "Conducted emission level [dBμV]; radiated emission level [dBμV/m]; "
                    "ESD, EFT, surge immunity results (Pass / Fail)"
                ),
                acceptance=(
                    "All conducted and radiated emission levels below IEC 61000-6-4 Class A limits. "
                    "All immunity tests pass at Level 3 per IEC 61000-6-2."
                ),
                method="T",
            ),
            TestSheet(
                ts_id="TS-6-06",
                title="Vibration Resistance / Mobile Applicability",
                objective=(
                    "Verify that the charger can withstand transportation vibrations and "
                    "is suitable for mobile applications as required by the project specification."
                ),
                standard="IEC 60068-2-6 (vibration sinusoidal) or IEC 60068-2-64 (random vibration); project specification (mobile-capable)",
                equipment=(
                    "Vibration test table or shake table (if available); "
                    "alternatively: perform a road-transport simulation or review based on module datasheet vibration rating. "
                    "Torque wrench; visual inspection checklist."
                ),
                setup=(
                    "Mount the charger (or representative assembly) on the vibration test fixture. "
                    "If no shake table is available, perform a structured visual and documentation review of the converter module vibration rating."
                ),
                procedure_steps=[
                    "Option A – Physical vibration test:",
                    "  Set vibration profile per IEC 60068-2-64 (road transport: 5–100 Hz, 0.1 g²/Hz).",
                    "  Run for 1 hour per axis (X, Y, Z).",
                    "  After test: inspect for loose fasteners, connector damage, coolant leaks.",
                    "Option B – Documentation review (if no shake table):",
                    "  Verify converter module datasheet vibration rating meets IEC 60068-2-6 / -64 road transport class.",
                    "  Verify all fasteners are torqued to specification.",
                    "  Verify all connectors are locked/latch-secured.",
                    "  Verify cooling hoses have strain-relief clamps.",
                    "Record outcome of physical test or documentation review.",
                ],
                measured_qty="Vibration test level [g or g²/Hz]; inspection result (Pass / Fail); any defects found",
                acceptance=(
                    "No structural damage, loose fasteners, connector failures or coolant leaks after vibration test. "
                    "Converter module vibration rating is suitable for road/mobile application."
                ),
                method="T / I",
            ),
            TestSheet(
                ts_id="TS-6-07",
                title="DC Leakage Current vs. e-PU10 Precharge Capacity",
                objective=(
                    "Verify that the DC leakage current of the charger is lower than the "
                    "e-PU10 battery precharge circuit capacity to prevent precharge failure."
                ),
                standard="e-PU10 restrictions (project plan); IEC 61557-8",
                equipment=(
                    "Calibrated insulation / leakage current meter (≤ 12 months calibration); "
                    "e-PU10 battery system precharge specification document"
                ),
                setup=(
                    "Charger powered up with DC output disconnected from battery. "
                    "Measure leakage current from DC+ and DC− rails to chassis/earth."
                ),
                procedure_steps=[
                    "Power up the charger (AC input connected, DC output open or connected to load at rated voltage).",
                    "Measure DC leakage current from DC+ bus to chassis ground [mA].",
                    "Measure DC leakage current from DC− bus to chassis ground [mA].",
                    "Obtain the e-PU10 precharge circuit maximum current rating from the e-PU10 specification.",
                    "Compare measured leakage against precharge circuit limit.",
                    "Record result.",
                ],
                measured_qty="DC leakage current I_leak_DC+ [mA]; I_leak_DC− [mA]",
                acceptance=(
                    "Both I_leak_DC+ and I_leak_DC− < e-PU10 precharge circuit current rating. "
                    "Specific limit to be confirmed from e-PU10 V2 development spec."
                ),
                method="T",
            ),
            TestSheet(
                ts_id="TS-6-08",
                title="Reliability / Uptime — 72-Hour Endurance Run",
                objective=(
                    "Verify that the charger achieves an operational uptime ≥ 98 % "
                    "during a representative endurance run."
                ),
                standard="Project specification (reliability > 98 % uptime); IEC 61439",
                equipment=(
                    "Automated test controller (to cycle load / charger commands); "
                    "event logger or data acquisition system; "
                    "AC power supply; DC electronic load"
                ),
                setup=(
                    "Connect charger to rated AC supply and DC electronic load. "
                    "Program automated cycling: 30 min full load → 5 min 50 % load → repeat. "
                    "Log all faults, shutdowns and restart events automatically."
                ),
                procedure_steps=[
                    "Start automated endurance cycle. Record start time.",
                    "Run continuously for 72 hours (minimum).",
                    "Log all fault events: timestamp, fault code, duration of downtime.",
                    "After 72 hours, stop the test and record stop time.",
                    "Calculate total fault downtime [min].",
                    "Calculate uptime %: (72 × 60 − Σ downtime_min) / (72 × 60) × 100.",
                    "Review fault log for recurring patterns.",
                ],
                measured_qty="Total test duration [h]; total downtime [min]; uptime [%]",
                acceptance="Uptime ≥ 98 % over the 72-hour endurance run. No unrecoverable fault.",
                method="T",
            ),
        ],
    },
    # -----------------------------------------------------------------------
    # TRL 8 — Sales Readiness
    # -----------------------------------------------------------------------
    {
        "number": "8",
        "title": "Sales Readiness",
        "weeks": "Week 17–18",
        "effort": "8 h",
        "objective": (
            "Prepare and validate the commercial-facing documents (technical datasheet and "
            "sales one-pager) for the 400 A charger, ensuring all stated specifications are "
            "supported by verified test results from TRL 6."
        ),
        "entry_criteria": [
            "TRL 6 tollgate is signed off.",
            "All TRL 6 test sheets have been completed and signed.",
        ],
        "exit_criteria": [
            "Technical datasheet is reviewed and figures match test results.",
            "Sales one-pager is approved by project lead.",
            "Tollgate TRL 8 review is signed off.",
        ],
        "sheets": [
            TestSheet(
                ts_id="TS-8-01",
                title="Technical Datasheet Figure Verification",
                objective=(
                    "Verify that all performance figures stated in the technical datasheet "
                    "are consistent with the measured results from TRL 6 test sheets."
                ),
                standard="Internal document control; ISO 9001 (document accuracy)",
                equipment="Technical datasheet draft; TRL 6 test results (TS-6-01 through TS-6-08)",
                setup="Compare each technical parameter in the datasheet against the corresponding TRL 6 measured result.",
                procedure_steps=[
                    "Obtain draft technical datasheet.",
                    "For each stated parameter, identify the corresponding TRL 6 test sheet result.",
                    "Check: (a) Rated power [kW]; (b) Efficiency [%]; (c) DC voltage range [V]; (d) Weight [kg]; (e) Dimensions [mm]; (f) Cooling specifications; (g) Communication protocol.",
                    "Verify that each datasheet figure is ≤ (for max limits) or ≥ (for min limits) the measured value plus tolerance.",
                    "Mark each parameter as Consistent / Inconsistent.",
                    "Correct any inconsistencies in the datasheet.",
                ],
                measured_qty="Number of parameters consistent with test results (count / 7)",
                acceptance="All 7 key parameters in the datasheet are consistent with measured TRL 6 results. No overstated figures.",
                method="I / A",
            ),
            TestSheet(
                ts_id="TS-8-02",
                title="Sales One-Pager Specification Sign-Off",
                objective=(
                    "Obtain formal approval of the sales one-pager document by the project "
                    "lead and at least one sales stakeholder."
                ),
                standard="Internal document approval process",
                equipment="Sales one-pager draft; sign-off sheet",
                setup="Present the final one-pager to project lead (L. Rietkerk) and a sales representative.",
                procedure_steps=[
                    "Distribute draft sales one-pager to project lead and sales representative.",
                    "Review document for: (a) accuracy of technical specs; (b) USPs clearly stated; (c) compliance claims supported by certificates.",
                    "Incorporate any review comments.",
                    "Obtain signed approval from project lead.",
                    "Obtain signed approval from sales representative.",
                    "File approved document in project document management system.",
                ],
                measured_qty="Number of approvals obtained (count / 2)",
                acceptance="Both project lead and sales representative have signed the one-pager. All comments resolved.",
                method="I",
            ),
        ],
    },
    # -----------------------------------------------------------------------
    # TRL 9 — Product Release
    # -----------------------------------------------------------------------
    {
        "number": "9",
        "title": "Product Release",
        "weeks": "Week 19 – 02-11-2026",
        "effort": "8 h",
        "objective": (
            "Finalise all product documentation, verify series-readiness of the design, "
            "and formally transfer the 400 A charger product to the sales matrix."
        ),
        "entry_criteria": [
            "TRL 8 tollgate is signed off.",
            "All test sheets (TRL 2 through TRL 8) are complete and signed.",
            "CE documentation package is complete.",
        ],
        "exit_criteria": [
            "Documentation completeness check is passed.",
            "Series-ready verification is signed off.",
            "Product is entered in the sales matrix.",
            "Final tollgate TRL 9 is signed off (deadline: 02-11-2026).",
        ],
        "sheets": [
            TestSheet(
                ts_id="TS-9-01",
                title="Final Documentation Completeness Check",
                objective=(
                    "Verify that all required product documents are complete, approved, "
                    "version-controlled and archived."
                ),
                standard="IEC 61439; NEN 1010; CE marking directive; internal document management",
                equipment="Document management system; document checklist",
                setup="Retrieve the complete document package from the project folder.",
                procedure_steps=[
                    "Check availability and approval status of each required document:",
                    "  a. Project plan (final version, signed)",
                    "  b. This Measurement Plan (all test sheets completed and signed)",
                    "  c. Detailed electrical drawings (released revision)",
                    "  d. Detailed mechanical drawings (released revision)",
                    "  e. BOM (final, with part numbers and supplier references)",
                    "  f. Software documentation (architecture, register map, version)",
                    "  g. CE Declaration of Conformity",
                    "  h. IEC 61439 design verification record",
                    "  i. NEN 1010 compliance record",
                    "  j. User manual / operating instructions",
                    "  k. Assembly / build book",
                    "Mark each document as Complete / Incomplete / Missing.",
                    "Raise non-conformances for any incomplete or missing items.",
                ],
                measured_qty="Number of documents Complete (count / 11)",
                acceptance="All 11 required documents are Complete, approved and version-controlled. Zero missing items.",
                method="I",
            ),
            TestSheet(
                ts_id="TS-9-02",
                title="Series-Ready Verification",
                objective=(
                    "Verify that the design is reproducible and manufacturable in series "
                    "production without further design changes."
                ),
                standard="Internal series-readiness criteria; IEC 61439",
                equipment="Series-readiness checklist; assembly instructions; BOM",
                setup="Review the complete design package against the series-readiness checklist.",
                procedure_steps=[
                    "Verify all components in the BOM are commercially available with lead time < 8 weeks.",
                    "Verify the assembly process can be performed by a trained technician using the build book alone (no expert assistance required).",
                    "Verify all software can be flashed/configured from the build book procedure.",
                    "Verify no prototype-specific workarounds remain in the design or build procedure.",
                    "Verify quality control steps are defined for production.",
                    "Record any open items.",
                ],
                measured_qty="Number of series-readiness criteria met (count / 5); number of open items",
                acceptance="All 5 criteria met. Zero open items. Design is cleared for series production.",
                method="I / A",
            ),
            TestSheet(
                ts_id="TS-9-03",
                title="Transfer to Sales Matrix Sign-Off",
                objective=(
                    "Formally transfer the validated 400 A charger product to the sales matrix "
                    "and confirm that all product information is correctly entered."
                ),
                standard="Internal sales and product management process",
                equipment="Sales matrix template; approved datasheet; pricing information",
                setup="Enter product into the sales matrix with the project lead and sales stakeholder present.",
                procedure_steps=[
                    "Enter product in sales matrix with the following information:",
                    "  a. Product name: 400A AC/DC Charger for e-PU10",
                    "  b. Rated power [kW]",
                    "  c. DC voltage range [V]",
                    "  d. Efficiency [%]",
                    "  e. Communication protocol",
                    "  f. Compliance standards (IEC 61439, NEN 1010, CE)",
                    "  g. Lead time and price indication",
                    "Obtain sign-off from project lead (L. Rietkerk).",
                    "Obtain sign-off from sales representative.",
                    "Archive the signed transfer document.",
                ],
                measured_qty="Number of sales matrix fields completed (count / 7); number of signatures obtained (count / 2)",
                acceptance="All 7 sales matrix fields are completed. Both signatures obtained. Transfer document archived.",
                method="I",
            ),
        ],
    },
]


# ---------------------------------------------------------------------------
# Document generation
# ---------------------------------------------------------------------------

def _add_cover(doc):
    """Add cover/title page content."""
    doc.add_paragraph()
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("400A AC/DC Charger")
    run.bold = True
    run.font.size = Pt(36)

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("Measurement Plan")
    run.bold = True
    run.font.size = Pt(28)

    doc.add_paragraph()
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("Requirements Verification Plan")
    run.font.size = Pt(16)

    doc.add_paragraph()
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("Based on: ProjectPlan_400A_Charger.docx")

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("Version: 1.0")

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run(f"Date: {date.today().isoformat()}")

    doc.add_paragraph()
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("Project Lead: Lennard Rietkerk")

    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run("Kickoff: 23-06-2026  |  Target: 02-11-2026")


def _add_introduction(doc):
    """Introduction section (mirrors the e-PU Cabinet V2 layout)."""
    doc.add_heading("Introduction", level=1)
    doc.add_paragraph(
        "This document is the Measurement Plan for the 400A AC/DC Charger project "
        "— an unidirectional, liquid-cooled AC/DC battery charger with galvanic isolation, "
        "designed for charging the main battery (600–800 VDC) of the e-PU10 BESS."
    )
    doc.add_paragraph(
        "Purpose: to make sure every requirement from the project plan is verified during "
        "the test phase, with nothing missed."
    )
    doc.add_paragraph(
        "This plan follows the standard verification approach (Verification Cross-Reference "
        "Matrix + test sheets). It is split into two easy-to-use parts:"
    )
    for item in [
        "Section 3 — Verification Cross-Reference Matrix: a short, at-a-glance list of every "
        "requirement with its method and status. Use it to see overall progress.",
        "Section 4 — Test Sheets: one small card per requirement with the exact test steps, "
        "pass criteria and equipment. Use it while performing each test.",
    ]:
        p = doc.add_paragraph(style="List Paragraph")
        p.add_run(item)

    doc.add_heading("How to use this plan", level=2)
    for step in [
        "Work through the Test Sheets (Section 4), one requirement at a time.",
        'For each sheet, follow "How to test", then write the measured value in "Result".',
        "Set Status to Pass or Fail based on the acceptance criterion.",
        "Copy the Pass/Fail result back into the matrix (Section 3) to track overall progress.",
        "Complete the Sign-off (Section 6) when all test sheets for a TRL phase are finished.",
    ]:
        p = doc.add_paragraph(style="List Paragraph")
        p.add_run(step)


def _add_legend(doc):
    """Legend section with verification methods and status values tables."""
    doc.add_heading("Legend", level=1)
    doc.add_heading("Verification Methods", level=2)
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0]
    for i, h in enumerate(["Code", "Method", "What it means"]):
        hdr.cells[i].text = h
    _make_row_header(hdr)
    data = [
        ("T", "Test", "Physically operate and measure with instruments."),
        ("I", "Inspection", "Visual check, dimensional check, certificate review."),
        ("A", "Analysis", "Calculation, simulation, or engineering judgement."),
        ("D", "Demonstration", "Show it works under normal operation, without measuring."),
    ]
    for code, method, desc in data:
        row = tbl.add_row()
        row.cells[0].text = code
        row.cells[1].text = method
        row.cells[2].text = desc

    doc.add_heading("Status Values", level=2)
    tbl2 = doc.add_table(rows=1, cols=2)
    tbl2.style = "Table Grid"
    hdr2 = tbl2.rows[0]
    hdr2.cells[0].text = "Status"
    hdr2.cells[1].text = "Meaning"
    _make_row_header(hdr2)
    for status, meaning in [
        ("Not started", "Test not performed yet."),
        ("In progress", "Test currently being executed."),
        ("Pass", "Verified — acceptance criterion met."),
        ("Fail", "Acceptance criterion not met — corrective action required."),
    ]:
        row = tbl2.add_row()
        row.cells[0].text = status
        row.cells[1].text = meaning


def _add_vcrm(doc, trl_data):
    """Verification Cross-Reference Matrix — at-a-glance list of every requirement.

    Mirrors the e-PU Cabinet V2 matrix: ID | Requirement | Method | Test stage | Status.
    """
    doc.add_heading("Verification Cross-Reference Matrix", level=1)
    doc.add_paragraph(
        "At-a-glance list of every requirement. Full test details are in Section 4 "
        "(Test Sheets). Fill in the Status column as testing progresses."
    )
    tbl = doc.add_table(rows=1, cols=5)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0]
    for i, h in enumerate(["ID", "Requirement", "Method", "Test stage", "Status"]):
        hdr.cells[i].text = h
    _make_row_header(hdr)

    for trl in trl_data:
        for sheet in trl["sheets"]:
            row = tbl.add_row()
            row.cells[0].text = sheet.ts_id
            row.cells[1].text = sheet.title
            row.cells[2].text = sheet.method
            row.cells[3].text = f"TRL {trl['number']}"
            row.cells[4].text = ""  # blank for filling in


def _add_test_sheet(doc, sheet: TestSheet, test_stage: str):
    """Render a single test sheet as a 2-column table (label | value).

    Uses the identical row structure as the e-PU Cabinet V2 test sheets:
    Requirement, Source, Verification method, How to test, Acceptance criterion,
    Equipment / tooling, Test stage, Result, Status.
    """
    tbl = doc.add_table(rows=0, cols=2)
    tbl.style = "Table Grid"
    # Set column widths roughly 30 / 70
    tbl.columns[0].width = Inches(2.1)
    tbl.columns[1].width = Inches(4.4)

    # Build the "How to test" cell from the setup + numbered procedure steps so
    # that no procedural detail is lost while keeping the e-PU row layout.
    how_to_test = sheet.setup.strip()
    if sheet.procedure_steps:
        steps = "\n".join(f"{i + 1}. {s}" for i, s in enumerate(sheet.procedure_steps))
        how_to_test = (how_to_test + "\n" if how_to_test else "") + steps

    acceptance = sheet.acceptance
    if sheet.measured_qty:
        acceptance = f"{acceptance}\nMeasured quantity: {sheet.measured_qty}"

    rows_data = [
        ("Requirement", sheet.objective),
        ("Source", sheet.standard),
        ("Verification method", sheet.method),
        ("How to test", how_to_test),
        ("Acceptance criterion (pass condition)", acceptance),
        ("Equipment / tooling", sheet.equipment),
        ("Test stage", test_stage),
        ("Result (fill in during testing)", ""),
        ("Status (Pass / Fail)", ""),
    ]

    for label, value in rows_data:
        row = tbl.add_row()
        row.cells[0].text = label
        row.cells[1].text = value
        _make_label_cell(row.cells[0])

    doc.add_paragraph()  # spacing between sheets


def _add_test_sheets(doc, trl_data):
    """Add the Test Sheets chapter: one TRL sub-section per phase (e-PU layout)."""
    doc.add_heading("Test Sheets", level=1)
    doc.add_paragraph(
        'One card per requirement. Follow "How to test", record the "Result", '
        "then mark Pass or Fail."
    )
    for trl in trl_data:
        doc.add_heading(f"TRL {trl['number']} — {trl['title']}", level=2)
        test_stage = f"TRL {trl['number']}"
        for sheet in trl["sheets"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{sheet.ts_id}  —  {sheet.title}")
            run.bold = True
            _add_test_sheet(doc, sheet, test_stage)


def _add_test_phase_overview(doc):
    """Summary table of TRL phases and what is tested (e-PU Cabinet V2 layout)."""
    doc.add_heading("Test Phase Overview", level=1)
    doc.add_paragraph(
        "Testing follows the TRL (Technology Readiness Level) phases from the project plan. "
        "Each stage builds on the previous one and ends with a formal tollgate review."
    )
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0]
    for i, h in enumerate(["Stage", "TRL", "What is tested", "Requirement IDs"]):
        hdr.cells[i].text = h
    _make_row_header(hdr)
    for i, trl in enumerate(TRL_DATA):
        row = tbl.add_row()
        row.cells[0].text = f"{i + 1} — {trl['title']}"
        row.cells[1].text = f"TRL {trl['number']}"
        row.cells[2].text = trl["objective"]
        row.cells[3].text = "; ".join(s.ts_id for s in trl["sheets"])


def _add_signoff(doc):
    """Sign-off table."""
    doc.add_heading("Sign-Off", level=1)
    doc.add_paragraph(
        "Complete this table when testing for a TRL phase or the full campaign is finished."
    )
    doc.add_paragraph(
        "Declaration: all requirements in this plan have been verified. "
        "Results are recorded in the Test Sheets above."
    )
    tbl = doc.add_table(rows=1, cols=4)
    tbl.style = "Table Grid"
    hdr = tbl.rows[0]
    for i, h in enumerate(["Role", "Name", "Signature", "Date"]):
        hdr.cells[i].text = h
    _make_row_header(hdr)
    for role in [
        "Tester / Test Engineer",
        "Test Lead / Project Lead",
        "Quality / Safety Engineer",
        "Approver",
    ]:
        row = tbl.add_row()
        row.cells[0].text = role
        for j in range(1, 4):
            row.cells[j].text = ""


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

def main():
    print(f"Opening template: {TEMPLATE_PATH}")
    doc = Document(TEMPLATE_PATH)

    # --- Clear all body content, keep sectPr ---
    from docx.oxml.ns import qn as _qn
    body = doc.element.body
    sect_pr = body.find(_qn("w:sectPr"))
    # Remove everything except sectPr
    to_remove = [child for child in list(body) if child.tag != _qn("w:sectPr")]
    for el in to_remove:
        body.remove(el)

    # --- Build new content (mirrors the e-PU Cabinet V2 chapter structure) ---
    _add_cover(doc)
    doc.add_page_break()

    _add_introduction(doc)
    doc.add_page_break()

    _add_legend(doc)
    doc.add_page_break()

    _add_vcrm(doc, TRL_DATA)
    doc.add_page_break()

    _add_test_sheets(doc, TRL_DATA)
    doc.add_page_break()

    _add_test_phase_overview(doc)
    doc.add_page_break()

    _add_signoff(doc)

    # --- Save ---
    os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)
    doc.save(OUTPUT_PATH)
    print(f"Saved: {OUTPUT_PATH}")


if __name__ == "__main__":
    main()
